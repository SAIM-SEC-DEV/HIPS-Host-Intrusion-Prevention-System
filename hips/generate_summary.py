import os
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx is not installed. Please install it using: pip install python-docx")
    exit(1)

def create_summary_docx():
    doc = Document()
    
    # Change default font to Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    summary_run = title.add_run('Project Summary\n\n')
    summary_run.bold = True
    summary_run.font.size = Pt(20)
    
    title_run = title.add_run('Project Title: HIPS-OOP (Enterprise-Grade Host Intrusion Prevention System)\n')
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle_run = title.add_run('Najam Sherazi\nApril 27, 2026\nCybersecurity Course Project, Islamabad')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(12)

    # Industry Problem
    heading1 = doc.add_heading('Industry Problem', level=1)
    
    p1 = doc.add_paragraph(
        "Modern enterprise networks face sophisticated, rapidly evolving cyber threats—such as zero-day malware, "
        "fileless attacks, and ransomware—that easily bypass traditional signature-based antivirus solutions. "
        "Cybercriminals continually adapt their tactics, exploiting systemic vulnerabilities to gain unauthorized access, "
        "exfiltrate sensitive data, and disrupt critical operations. According to recent industry statistics, "
        "the average cost of a data breach has reached record highs, and the time required to identify and contain a "
        "breach often spans hundreds of days."
    )
    p2 = doc.add_paragraph(
        "A major gap exists in deploying lightweight, intelligent host-based systems that can monitor deep system behaviors "
        "(such as unauthorized registry modifications, unexpected network ports opening, and abnormal process execution). "
        "Many existing solutions are resource-heavy, generate excessive false positives, or fail to map anomalies "
        "to standardized threat models. There is a critical need for a proactive endpoint security solution that detects "
        "attacks dynamically, identifies threat tactics, and isolates compromised hosts in real-time to prevent lateral movement."
    )

    # Solution
    heading2 = doc.add_heading('Our Solution: HIPS-OOP Enterprise Edition', level=1)
    
    p3 = doc.add_paragraph(
        "To address these critical security gaps, we developed the HIPS-OOP Enterprise Edition—a robust, modular, "
        "and service-based Host Intrusion Prevention System. This project is designed to continuously monitor, detect, "
        "and respond to malicious endpoint activities through an intelligent agent architecture and a centralized management backend."
    )

    p4 = doc.add_paragraph("Key Implementation Details and Methods:")
    
    # Bullet points
    doc.add_paragraph("Decoupled Agent Architecture: Utilizes a Service Manager pattern to independently handle File, Network, Process, Registry, Memory, and Asset monitoring modules, eliminating single points of failure.", style='List Bullet')
    doc.add_paragraph("Real-Time Behavioral Auditing: Continuously tracks operations, flagging anomalies based on intelligent baseline comparisons and alerting administrators to suspicious file modifications and unauthorized registry edits.", style='List Bullet')
    doc.add_paragraph("MITRE ATT&CK Integration: Events are mapped directly to the industry-standard MITRE ATT&CK framework (e.g., matching tactics and technique IDs), aiding analysts in understanding threat progression.", style='List Bullet')
    doc.add_paragraph("Automated Response Capabilities: Features a server-to-agent command queue allowing immediate actions such as blocking IPs, killing malicious processes, or quarantining affected systems.", style='List Bullet')
    doc.add_paragraph("Enterprise Security Hardening: Incorporates strict input sanitization to prevent command injection, robust API authentication via tokens, password policies, and log isolation mechanisms.", style='List Bullet')
    doc.add_paragraph("Centralized Management Dashboard: A secure, PHP/MySQL-backed web interface that provides live alerting, device status, detailed threat reporting, and remote command orchestration.", style='List Bullet')

    doc.add_paragraph()
    p5 = doc.add_paragraph("Conclusion:")
    p5_run = p5.add_run("By combining proactive endpoint telemetry with an extensible, secure backend, HIPS-OOP significantly reduces the dwell time of advanced threats, transforming endpoint defense into a highly dynamic and resilient operation.")
    p5_run.italic = True

    # Save document
    filename = 'HIPS_Project_Summary.docx'
    doc.save(filename)
    print(f"Successfully generated {filename}")

if __name__ == '__main__':
    create_summary_docx()
